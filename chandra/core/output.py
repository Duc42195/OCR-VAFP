"""Output processing and formatting."""
import hashlib
import json
import re
import math
from dataclasses import dataclass, asdict
from functools import lru_cache
from typing import Tuple

import six
from PIL import Image
from bs4 import BeautifulSoup
from markdownify import MarkdownConverter, re_whitespace

from chandra.settings import settings


@lru_cache
def _hash_html(html: str):
    """Hash HTML content for generating unique filenames."""
    return hashlib.md5(html.encode("utf-8")).hexdigest()


def get_image_name(html: str, div_idx: int):
    """Generate unique image filename from HTML and block index."""
    html_hash = _hash_html(html)
    return f"{html_hash}_{div_idx}_img.webp"


def extract_images(html: str, chunks: dict, image: Image.Image):
    """Extract images from layout blocks and save as separate files."""
    images = {}
    div_idx = 0
    for idx, chunk in enumerate(chunks):
        div_idx += 1
        if chunk["label"] in ["Image", "Figure"]:
            img = chunk["content"].find("img")
            if not img:
                continue
            bbox = chunk["bbox"]
            try:
                block_image = image.crop(bbox)
            except ValueError:
                # Happens when bbox coordinates are invalid
                continue
            img_name = get_image_name(html, div_idx)
            images[img_name] = block_image
    return images


def parse_html(
    html: str, include_headers_footers: bool = False, include_images: bool = True
):
    """Parse and clean HTML output from model."""
    soup = BeautifulSoup(html, "html.parser")
    top_level_divs = soup.find_all("div", recursive=False)
    out_html = ""
    image_idx = 0
    div_idx = 0
    for div in top_level_divs:
        div_idx += 1
        label = div.get("data-label")

        # Skip headers and footers if not included
        if label and not include_headers_footers:
            if label in ["Page-Header", "Page-Footer"]:
                continue
        if label and not include_images:
            if label in ["Image", "Figure"]:
                continue

        if label in ["Image", "Figure"]:
            img = div.find("img")
            img_src = get_image_name(html, div_idx)

            # If no tag, add one in
            if img:
                img["src"] = img_src
                image_idx += 1
            else:
                img = BeautifulSoup(f"<img src='{img_src}'/>", "html.parser")
                div.append(img)

        # Wrap text content in <p> tags if no inner HTML tags exist
        if label in ["Text"] and not re.search(
            "<.+>", str(div.decode_contents()).strip()
        ):
            # Add inner p tags if missing for text blocks
            text_content = str(div.decode_contents()).strip()
            text_content = f"<p>{text_content}</p>"
            div.clear()
            div.append(BeautifulSoup(text_content, "html.parser"))

        content = str(div.decode_contents())
        out_html += content
    return out_html


class Markdownify(MarkdownConverter):
    """Custom Markdown converter with math and table support."""
    def __init__(
        self,
        inline_math_delimiters,
        block_math_delimiters,
        **kwargs,
    ):
        super().__init__(**kwargs)
        self.inline_math_delimiters = inline_math_delimiters
        self.block_math_delimiters = block_math_delimiters

    def convert_math(self, el, text, parent_tags):
        """Convert math tags to LaTeX-compatible format."""
        block = el.has_attr("display") and el["display"] == "block"
        if block:
            return (
                "\n"
                + self.block_math_delimiters[0]
                + text.strip()
                + self.block_math_delimiters[1]
                + "\n"
            )
        else:
            return (
                " "
                + self.inline_math_delimiters[0]
                + text.strip()
                + self.inline_math_delimiters[1]
                + " "
            )

    def convert_table(self, el, text, parent_tags):
        """Keep tables as HTML."""
        return "\n\n" + str(el) + "\n\n"

    def convert_a(self, el, text, parent_tags):
        """Convert links with proper escaping."""
        text = self.escape(text)
        # Escape brackets and parentheses in text
        text = re.sub(r"([\[\]()])", r"\\\1", text)
        return super().convert_a(el, text, parent_tags)

    def escape(self, text, parent_tags=None):
        """Escape special characters in text."""
        text = super().escape(text, parent_tags)
        if self.options["escape_dollars"]:
            text = text.replace("$", r"\$")
        return text

    def process_text(self, el, parent_tags=None):
        """Process text nodes with proper normalization."""
        text = six.text_type(el) or ""

        # normalize whitespace if we're not inside a preformatted element
        if not el.find_parent("pre"):
            text = re_whitespace.sub(" ", text)

        # escape special characters if we're not inside a preformatted or code element
        if not el.find_parent(["pre", "code", "kbd", "samp", "math"]):
            text = self.escape(text)

        # remove trailing whitespaces if any of the following condition is true:
        # - current text node is the last node in li
        # - current text node is followed by an embedded list
        if el.parent.name == "li" and (
            not el.next_sibling or el.next_sibling.name in ["ul", "ol"]
        ):
            text = text.rstrip()

        return text


def parse_markdown(
    html: str, include_headers_footers: bool = False, include_images: bool = True
):
    """Convert HTML to Markdown format."""
    html = parse_html(html, include_headers_footers, include_images)

    md_cls = Markdownify(
        heading_style="ATX",
        bullets="-",
        escape_misc=False,
        escape_underscores=True,
        escape_asterisks=True,
        escape_dollars=True,
        sub_symbol="<sub>",
        sup_symbol="<sup>",
        inline_math_delimiters=("$", "$"),
        block_math_delimiters=("$$", "$$"),
    )
    try:
        markdown = md_cls.convert(html)
    except Exception as e:
        print(f"Error converting HTML to Markdown: {e}")
        markdown = ""
    return markdown.strip()

def fix_json_tail(s: str) -> str:
    """
    Fix truncated JSON where model stops early.
    Only appends missing ] or } at the end.
    """

    s = s.rstrip()

    # Đếm số ngoặc
    open_curly = s.count("{")
    close_curly = s.count("}")
    open_square = s.count("[")
    close_square = s.count("]")

    # Vá ngoặc vuông trước
    if close_square < open_square:
        s += "]" * (open_square - close_square)

    # Vá ngoặc nhọn sau
    if close_curly < open_curly:
        s += "}" * (open_curly - close_curly)

    return s


def extract_first_fenced_json_autofix(raw: str) -> dict:
    start = raw.find("{")
    if start == -1:
        raise ValueError("No JSON object start '{' found")

    json_str = raw[start:].strip()

    # 2. Remove trailing fences if any (``` or ```json)
    for fence in ["```json", "```"]:
        if fence in json_str:
            json_str = json_str.split(fence)[0].strip()

    # ---- AUTO FIX TAIL ----
    json_str = fix_json_tail(json_str)

    return json.loads(json_str)

def save_json(data, path="output.json"):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(
            data,
            f,
            ensure_ascii=False,
            indent=2
        )

@dataclass
class LayoutBlock:
    """Represents a layout block in the document."""
    bbox: list[int]
    label: str
    content: str


def parse_layout(html: str, image: Image.Image, bbox_scale=settings.BBOX_SCALE):
    """Parse layout blocks from HTML output."""
    soup = BeautifulSoup(html, "html.parser")
    top_level_divs = soup.find_all("div", recursive=False)
    width, height = image.size
    width_scaler = width / bbox_scale
    height_scaler = height / bbox_scale
    layout_blocks = []
    for div in top_level_divs:
        bbox = div.get("data-bbox")

        try:
            bbox = json.loads(bbox)
            assert len(bbox) == 4, "Invalid bbox length"
        except Exception:
            try:
                bbox = bbox.split(" ")
                assert len(bbox) == 4, "Invalid bbox length"
            except Exception:
                bbox = [0, 0, 1, 1]

        bbox = list(map(int, bbox))
        # Normalize bbox
        bbox = [
            max(0, int(bbox[0] * width_scaler)),
            max(0, int(bbox[1] * height_scaler)),
            min(int(bbox[2] * width_scaler), width),
            min(int(bbox[3] * height_scaler), height),
        ]
        label = div.get("data-label", "block")
        content = str(div.decode_contents())
        layout_blocks.append(LayoutBlock(bbox=bbox, label=label, content=content))
    return layout_blocks


def parse_chunks(html: str, image: Image.Image, bbox_scale=settings.BBOX_SCALE):
    """Parse layout blocks as chunks dict."""
    layout = parse_layout(html, image, bbox_scale=bbox_scale)
    chunks = [asdict(block) for block in layout]
    return chunks


def scale_to_fit(
    img: Image.Image,
    max_size: Tuple[int, int] = (3072, 2048),
    min_size: Tuple[int, int] = (28, 28),
):
    """Scale image to fit within size constraints."""
    resample_method = Image.Resampling.LANCZOS

    width, height = img.size

    # Check for empty or invalid image
    if width == 0 or height == 0:
        return img

    max_width, max_height = max_size
    min_width, min_height = min_size

    current_pixels = width * height
    max_pixels = max_width * max_height
    min_pixels = min_width * min_height

    if current_pixels > max_pixels:
        scale_factor = (max_pixels / current_pixels) ** 0.5

        new_width = math.floor(width * scale_factor)
        new_height = math.floor(height * scale_factor)
    elif current_pixels < min_pixels:
        scale_factor = (min_pixels / current_pixels) ** 0.5

        new_width = math.ceil(width * scale_factor)
        new_height = math.ceil(height * scale_factor)
    else:
        return img

    return img.resize((new_width, new_height), resample=resample_method)


def detect_repeat_token(
    predicted_tokens: str,
    base_max_repeats: int = 4,
    window_size: int = 500,
    cut_from_end: int = 0,
    scaling_factor: float = 3.0,
):
    """Detect if output has repeated tokens indicating generation failure."""
    try:
        predicted_tokens = parse_markdown(predicted_tokens)
    except Exception as e:
        print(f"Error parsing markdown: {e}")
        return True

    if cut_from_end > 0:
        predicted_tokens = predicted_tokens[:-cut_from_end]

    for seq_len in range(1, window_size // 2 + 1):
        # Extract the potential repeating sequence from the end
        candidate_seq = predicted_tokens[-seq_len:]

        # Inverse scaling: shorter sequences need more repeats
        max_repeats = int(base_max_repeats * (1 + scaling_factor / seq_len))

        # Count how many times this sequence appears consecutively at the end
        repeat_count = 0
        pos = len(predicted_tokens) - seq_len
        if pos < 0:
            continue

        while pos >= 0:
            if predicted_tokens[pos : pos + seq_len] == candidate_seq:
                repeat_count += 1
                pos -= seq_len
            else:
                break

        if repeat_count > max_repeats:
            return True

    return False


def draw_layout(image: Image.Image, layout_blocks: list[LayoutBlock]):
    """Draw layout blocks with bounding boxes on image."""
    from PIL.ImageDraw import ImageDraw
    draw_image = image.copy()
    draw = ImageDraw(draw_image)
    for block in layout_blocks:
        if block.bbox[2] <= block.bbox[0] or block.bbox[3] <= block.bbox[1]:
            continue

        draw.rectangle(block.bbox, outline="red", width=2)
        draw.text((block.bbox[0], block.bbox[1]), block.label, fill="blue")

    return draw_image


# ============= QR Code Generation =============

def generate_viet_qr(
    amount: str,
    bank_number: str = "1234567890",
    purpose: str = "Thanh toan don hang",
    output_path: str = "viet_qr.png"
) -> Image.Image:
    """Generate Vietnamese QR Pay (Viet QR) code.
    
    Args:
        amount: Payment amount in VND
        bank_number: Bank account number
        purpose: Payment purpose
        output_path: Path to save QR image
        
    Returns:
        PIL Image of the QR code
    """
    try:
        from vietnam_qr_pay import QRPay, BanksObject
        import qrcode
        
        qr_pay = QRPay.init_viet_qr(
            bank_bin=BanksObject["vietcombank"].bin,
            bank_number=bank_number,
            amount=amount,
            purpose=purpose
        )
        content = qr_pay.build()
        qr_img = qrcode.make(content)
        qr_img.save(output_path)
        return qr_img
    except ImportError:
        print("Warning: vietnam_qr_pay or qrcode not installed. Skipping QR generation.")
        return None
    except Exception as e:
        print(f"Error generating Viet QR: {e}")
        return None


def generate_qr(
    data: str,
    output_path: str = "qr_code.png"
) -> Image.Image:
    """Generate standard QR code from data.
    
    Args:
        data: Data to encode in QR code
        output_path: Path to save QR image
        
    Returns:
        PIL Image of the QR code
    """
    try:
        import qrcode
        qr_img = qrcode.make(data)
        qr_img.save(output_path)
        return qr_img
    except ImportError:
        print("Warning: qrcode not installed. Skipping QR generation.")
        return None
    except Exception as e:
        print(f"Error generating QR code: {e}")
        return None

# ============= Extract Metadata (Title, Subtitle) =============
def extract_metadata_from_json(json_data):
    """Extract page title and subtitle from JSON."""
    metadata = {
        "title": "",
        "subtitle": ""
    }
    
    if "page" in json_data:
        page = json_data["page"]
        metadata["title"] = page.get("title", "")
        metadata["subtitle"] = page.get("subtitle", "")
    
    return metadata


def update_metadata_in_json(json_data, metadata):
    """Update page title and subtitle in JSON."""
    if "page" in json_data:
        json_data["page"]["title"] = metadata.get("title", "")
        json_data["page"]["subtitle"] = metadata.get("subtitle", "")
    return json_data


# ============= Edit json =============
def extract_tables_from_json(json_data):
    """Extract all tables from JSON, preserving their structure."""
    tables = []

    def walk(node, parent_ref=None, key=None):
        if isinstance(node, dict):
            if node.get("type") == "table" and "rows" in node:
                # Store reference to parent and key for updating
                tables.append({
                    "data": node,
                    "parent_ref": parent_ref,
                    "key": key
                })

            for k, v in node.items():
                walk(v, node, k)

        elif isinstance(node, list):
            for i, item in enumerate(node):
                walk(item, node, i)

    walk(json_data)
    return tables


def table_block_to_df(table_block):
    """Convert table block to pandas DataFrame."""
    rows = table_block.get("rows", [])

    clean_rows = []
    for r in rows:
        if isinstance(r, list):
            clean_rows.append(r)

    if not clean_rows:
        return pd.DataFrame()

    max_cols = max(len(r) for r in clean_rows)
    normalized = [r + [""] * (max_cols - len(r)) for r in clean_rows]

    return pd.DataFrame(normalized)


def df_to_table_block(df, table_block):
    """Convert DataFrame back to table block format."""
    table_block["rows"] = [list(df.columns)] + df.astype(str).values.tolist()


# ============= Save file =============
def export_json(json_data: dict, base_name: str):
    data = json.dumps(json_data, ensure_ascii=False, indent=2).encode("utf-8")
    filename = f"{base_name}.json"
    return data, "application/json", filename

from io import BytesIO
import pandas as pd
from openpyxl import load_workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font

def export_xlsx(json_data: dict, base_name: str):
    output = BytesIO()

    # Tạo workbook rỗng
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        # Tạo 1 sheet duy nhất
        sheet_name = "OCR_Result"
        pd.DataFrame().to_excel(writer, sheet_name=sheet_name, index=False)
        ws = writer.book[sheet_name]

        current_row = 1
        
        # ---- Add metadata (title, subtitle) ----
        metadata = extract_metadata_from_json(json_data)
        
        if metadata.get("title"):
            title_cell = ws.cell(row=current_row, column=1)
            title_cell.value = metadata.get("title")
            title_cell.font = Font(bold=True, size=12)
            current_row += 1
        
        if metadata.get("subtitle"):
            subtitle_cell = ws.cell(row=current_row, column=1)
            subtitle_cell.value = metadata.get("subtitle")
            current_row += 1
        
        current_row += 1  # Space after metadata

        tables = extract_tables_from_json(json_data)

        for i, table_item in enumerate(tables):
            # Extract table from wrapped dict
            if isinstance(table_item, dict) and "data" in table_item:
                table = table_item["data"]
            else:
                table = table_item
            df = table_block_to_df(table)

            if df.empty:
                continue

            # ---- Table title ----
            title_cell = ws.cell(row=current_row, column=1)
            title_cell.value = f"Table {i + 1}"
            title_cell.font = Font(bold=True)

            current_row += 1

            # ---- Write dataframe (WITHOUT index) ----
            for r in dataframe_to_rows(df, index=False, header=True):
                ws.append(r)
                current_row += 1

            # ---- Space between tables ----
            current_row += 2

    output.seek(0)

    return (
        output.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        f"{base_name}.xlsx",
    )


from docx import Document
from io import BytesIO

def export_docx(json_data: dict, base_name: str):
    from docx.shared import Pt
    
    doc = Document()
    
    # ---- Add metadata (title, subtitle) ----
    metadata = extract_metadata_from_json(json_data)
    
    if metadata.get("title"):
        title_para = doc.add_paragraph(metadata.get("title"))
        title_para.runs[0].font.size = Pt(14)
        title_para.runs[0].font.bold = True
    
    if metadata.get("subtitle"):
        subtitle_para = doc.add_paragraph(metadata.get("subtitle"))
        subtitle_para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Space after metadata
    
    tables = extract_tables_from_json(json_data)

    for table_item in tables:
        # Extract table from wrapped dict
        if isinstance(table_item, dict) and "data" in table_item:
            table = table_item["data"]
        else:
            table = table_item
        df = table_block_to_df(table)
        # Create table: df.shape[0] rows + 1 header row, df.shape[1] columns
        t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])

        # Write header (WITHOUT index)
        for j, col in enumerate(df.columns):
            t.rows[0].cells[j].text = str(col)

        # Write data (WITHOUT index)
        for i, row in df.iterrows():
            for j, val in enumerate(row):
                t.rows[i+1].cells[j].text = str(val)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return (
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{base_name}.docx"
    )

# from reportlab.platypus import SimpleDocTemplate, Paragraph
# from reportlab.lib.styles import getSampleStyleSheet
# from io import BytesIO

# def export_pdf(json_data: dict, base_name: str):
#     buf = BytesIO()
#     doc = SimpleDocTemplate(buf)
#     styles = getSampleStyleSheet()
#     elements = []

#     tables = extract_tables_from_json(json_data)

#     for table in tables:
#         df = table_block_to_df(table)
#         elements.append(Paragraph(df.to_string(), styles["Normal"]))

#     doc.build(elements)
#     buf.seek(0)

#     return buf.getvalue(), "application/pdf", f"{base_name}.pdf"

def export_from_json(json_data, format, base_name):
    if format == "json":
        return export_json(json_data, base_name)
    if format == "xlsx":
        return export_xlsx(json_data, base_name)
    if format == "docx":
        return export_docx(json_data, base_name)
    else:
        raise ValueError("Unsupported format")

